from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


SAMPLE_DIR = Path(__file__).resolve().parent.parent / "data" / "sample_uploads"


def create_pdf() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    output_path = SAMPLE_DIR / "客户数据导出管理规范.pdf"
    document = canvas.Canvas(str(output_path), pagesize=A4)
    document.setTitle("客户数据导出管理规范")
    document.setFont("STSong-Light", 18)
    document.drawString(72, 780, "客户数据导出管理规范")
    document.setFont("STSong-Light", 11)
    lines = [
        "1. 完整客户名单导出前，必须取得区域经理审批。",
        "2. 导出文件只能用于已批准的业务目的，并保存不得超过 7 天。",
        "3. 文件应存放在受控位置，禁止通过个人邮箱或即时通信工具外发。",
        "4. 发现未授权导出时，应立即停止使用并通知数据保护负责人。",
    ]
    y_position = 735
    for line in lines:
        document.drawString(72, y_position, line)
        y_position -= 32
    document.save()


def create_docx() -> None:
    output_path = SAMPLE_DIR / "客户数据导出补充规范.docx"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("客户数据导出补充规范")
    title_run.bold = True
    title_run.font.size = Pt(18)

    document.add_paragraph("本规范用于补充客户名单导出的审批、留存和审计要求。")
    document.add_paragraph("导出完整客户名单前，业务申请人必须获得区域经理审批，并在工单中保留审批记录。")
    document.add_paragraph("导出文件的默认保存期限为 7 天，到期后应按流程删除。")

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["风险等级", "控制要求", "建议动作"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    records = [
        ("高", "未审批导出客户名单", "立即停止导出并补充审批记录"),
        ("中", "导出文件保存超过 7 天", "删除过期文件并保留处置记录"),
    ]
    for record in records:
        cells = table.add_row().cells
        for cell, value in zip(cells, record):
            cell.text = value
    document.save(output_path)


def main() -> None:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    create_pdf()
    create_docx()
    print(f"Created PDF and DOCX samples in {SAMPLE_DIR}")


if __name__ == "__main__":
    main()
