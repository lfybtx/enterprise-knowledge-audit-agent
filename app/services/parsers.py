from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any, Optional


SUPPORTED_EXTENSIONS = {".docx", ".htm", ".html", ".pdf", ".txt", ".xlsx"}


class UnsupportedFileTypeError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedSection:
    text: str
    location: dict[str, Any]


@dataclass(frozen=True)
class ParsedDocument:
    file_type: str
    sections: list[ParsedSection]

    @property
    def text(self) -> str:
        return "\n\n".join(
            f"{location_marker(section.location)}\n{section.text}" for section in self.sections
        )


def parse_document(filename: str, content: bytes) -> tuple[str, str]:
    parsed = parse_document_sections(filename, content)
    return parsed.file_type, parsed.text


def parse_document_sections(filename: str, content: bytes) -> ParsedDocument:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(f"Unsupported file type '{extension}'. Supported types: {supported}")

    if extension == ".txt":
        return ParsedDocument("txt", parse_txt_sections(content))
    if extension in {".html", ".htm"}:
        return ParsedDocument("html", parse_html_sections(content))
    if extension == ".pdf":
        return ParsedDocument("pdf", parse_pdf_sections(content))
    if extension == ".docx":
        return ParsedDocument("docx", parse_docx_sections(content))
    if extension == ".xlsx":
        return ParsedDocument("xlsx", parse_xlsx_sections(content))

    raise UnsupportedFileTypeError(f"Unsupported file type '{extension}'")


def parse_txt(content: bytes) -> str:
    return parse_document_sections("document.txt", content).text


def parse_txt_sections(content: bytes) -> list[ParsedSection]:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = content.decode("utf-8", errors="ignore")

    normalize_text(text)
    return [
        ParsedSection(line.strip(), {"kind": "lines", "start_line": index, "end_line": index})
        for index, line in enumerate(text.splitlines(), start=1)
        if line.strip()
    ]


def parse_html_sections(content: bytes) -> list[ParsedSection]:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            html = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        html = content.decode("utf-8", errors="ignore")

    extractor = HtmlTextExtractor()
    extractor.feed(html)
    extractor.close()
    sections = [
        ParsedSection(text, {"kind": "html_paragraph", "paragraph_number": index})
        for index, text in enumerate(extractor.sections(), start=1)
    ]
    normalize_text("\n\n".join(section.text for section in sections))
    return sections


class HtmlTextExtractor(HTMLParser):
    BLOCK_TAGS = {
        "article",
        "blockquote",
        "br",
        "div",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "li",
        "p",
        "section",
        "td",
        "th",
        "tr",
    }
    SKIP_TAGS = {"script", "style", "noscript", "svg"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._skip_depth = 0
        self._current: list[str] = []
        self._sections: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        if tag in self.SKIP_TAGS:
            self._skip_depth += 1
            return
        if tag in self.BLOCK_TAGS:
            self._flush()

    def handle_endtag(self, tag: str) -> None:
        if tag in self.SKIP_TAGS and self._skip_depth:
            self._skip_depth -= 1
            return
        if tag in self.BLOCK_TAGS:
            self._flush()

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        cleaned = " ".join(data.split())
        if cleaned:
            self._current.append(cleaned)

    def sections(self) -> list[str]:
        self._flush()
        return self._sections

    def _flush(self) -> None:
        text = " ".join(self._current).strip()
        self._current = []
        if text:
            self._sections.append(text)


def parse_pdf(content: bytes) -> str:
    return parse_document_sections("document.pdf", content).text


def parse_pdf_sections(content: bytes) -> list[ParsedSection]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("PDF parser is unavailable. Install pypdf first.") from exc

    try:
        reader = PdfReader(BytesIO(content))
        pages: list[ParsedSection] = []
        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(ParsedSection(clean_text(page_text), {"kind": "page", "page_number": page_number}))
    except Exception as exc:
        raise DocumentParseError("Unable to read this PDF. Please upload a valid, unencrypted PDF.") from exc

    if not pages:
        pages = parse_pdf_ocr_sections(content)
    normalize_text("\n\n".join(page.text for page in pages))
    return pages


def parse_pdf_ocr_sections(content: bytes) -> list[ParsedSection]:
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as exc:
        raise DocumentParseError(
            "PDF has no extractable text and OCR dependencies are unavailable. Install pdf2image, pytesseract, Poppler, and Tesseract."
        ) from exc

    try:
        images = convert_from_bytes(content, dpi=220, fmt="png", thread_count=1)
    except Exception as exc:
        raise DocumentParseError("Unable to render scanned PDF pages for OCR.") from exc

    sections: list[ParsedSection] = []
    for page_number, image in enumerate(images, start=1):
        try:
            page_text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        except Exception as exc:
            raise DocumentParseError("OCR failed while reading scanned PDF pages.") from exc
        cleaned = clean_text(page_text)
        if cleaned:
            sections.append(
                ParsedSection(cleaned, {"kind": "page", "page_number": page_number, "ocr": True})
            )

    if not sections:
        raise EmptyDocumentError("OCR completed but no readable text was found in this PDF.")
    return sections


def parse_docx(content: bytes) -> str:
    return parse_document_sections("document.docx", content).text


def parse_docx_sections(content: bytes) -> list[ParsedSection]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("Word parser is unavailable. Install python-docx first.") from exc

    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("Unable to read this Word document. Please upload a valid .docx file.") from exc

    sections = [
        ParsedSection(clean_text(paragraph.text), {"kind": "paragraph", "paragraph_number": paragraph_number})
        for paragraph_number, paragraph in enumerate(document.paragraphs, start=1)
        if paragraph.text.strip()
    ]
    for table_number, table in enumerate(document.tables, start=1):
        headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells] if table.rows else []
        for row_number, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                fields = [
                    f"{headers[index] if index < len(headers) and headers[index] else f'Column {index + 1}'}: {value}"
                    for index, value in enumerate(cells)
                    if value
                ]
                sections.append(
                    ParsedSection(
                        " | ".join(fields),
                        {"kind": "table_row", "table_number": table_number, "row_number": row_number},
                    )
                )
    normalize_text("\n\n".join(section.text for section in sections))
    return sections


def parse_xlsx(content: bytes) -> str:
    return parse_document_sections("document.xlsx", content).text


def parse_xlsx_sections(content: bytes) -> list[ParsedSection]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentParseError("Excel parser is unavailable. Install openpyxl first.") from exc

    try:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    except Exception as exc:
        raise DocumentParseError("Unable to read this Excel workbook. Please upload a valid .xlsx file.") from exc

    sections: list[ParsedSection] = []
    try:
        for worksheet in workbook.worksheets:
            rows = list(enumerate(worksheet.iter_rows(values_only=True), start=1))
            non_empty_rows = [(row_number, [cell_to_text(value) for value in row]) for row_number, row in rows if any(row)]
            if not non_empty_rows:
                continue

            _, headers = non_empty_rows[0]
            if len(non_empty_rows) == 1:
                sections.append(
                    ParsedSection(
                        " | ".join(headers),
                        {"kind": "sheet_row", "sheet_name": worksheet.title, "row_number": 1},
                    )
                )
            else:
                for row_number, row in non_empty_rows[1:]:
                    fields = []
                    for column_index, value in enumerate(row):
                        if not value:
                            continue
                        header = headers[column_index] if column_index < len(headers) and headers[column_index] else f"Column {column_index + 1}"
                        fields.append(f"{header}: {value}")
                    if fields:
                        sections.append(
                            ParsedSection(
                                " | ".join(fields),
                                {"kind": "sheet_row", "sheet_name": worksheet.title, "row_number": row_number},
                            )
                        )
    finally:
        workbook.close()

    normalize_text("\n\n".join(section.text for section in sections))
    return sections


def cell_to_text(value: object | None) -> str:
    if value is None:
        return ""
    return str(value).strip()


def normalize_text(text: str) -> str:
    normalized = clean_text(text)
    if len(normalized) < 20:
        raise EmptyDocumentError("Document text is too short after parsing")
    return normalized


def clean_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())


def location_marker(location: dict[str, Any]) -> str:
    kind = location["kind"]
    if kind == "page":
        return f"[Page {location['page_number']}]"
    if kind == "paragraph":
        return f"[Paragraph {location['paragraph_number']}]"
    if kind == "table_row":
        return f"[Table {location['table_number']} Row {location['row_number']}]"
    if kind == "sheet_row":
        return f"[Sheet: {location['sheet_name']}] [Row {location['row_number']}]"
    if kind == "html_paragraph":
        return f"[HTML Paragraph {location['paragraph_number']}]"
    return f"[Lines {location['start_line']}-{location['end_line']}]"
